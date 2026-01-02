# filepath: c:\Users\srame\Documents\Coding\AutomaticSummaryGenerator\CsvLoader.py
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor


# -------------------------------------------------
# 1. Načtení CSV – sloupec Date jako datetime
# -------------------------------------------------
df = pd.read_csv(
    '20260101-Export-Sramek.csv',
    sep=',',
    encoding='utf-8',
    parse_dates=['Date']          # ← automaticky převede na datetime
)
billingPeriodMonth = df['Date'].dt.month
HOURLY_RATE = 350  # Kč za hodinu
CZECH_MONTHS = {
    1:  'leden',
    2:  'únor',
    3:  'březen',
    4:  'duben',
    5:  'květen',
    6:  'červen',
    7:  'červenec',
    8:  'srpen',
    9:  'září',
    10: 'říjen',
    11: 'listopad',
    12: 'prosinec'
}

headingMonth = CZECH_MONTHS[billingPeriodMonth.iloc[0]]
# -------------------------------------------------
# 3. Vytvoření Word‑dokumentu a tabulky
# -------------------------------------------------
doc = Document()
# heading
headingText = f'Výkaz odpracované doby zaměstnancem u dohody o pracovní činnosti za měsíc {headingMonth} roku {df['Date'].dt.year.iloc[0]}'
heading = doc.add_heading(headingText, level=1)
runHeading = heading.runs[0]
runHeading.font.name = 'Arial'
runHeading.font.size = Pt(16)
runHeading.bold = True
runHeading.font.color.rgb = RGBColor(0, 0, 0)
# Intro
intro = doc.add_paragraph()
run = intro.add_run("Práce je po dohodě smluvních stran vykonávána na dálku. V souladu s § 87a zákoníku práce si bude zaměstnanec pracovní dobu do směn při práci na dálku rozvrhovat sám, práce nebude vykonávána ve dnech svátků, víkendů a v čase od 22:00 do 6:00 hod. Délka směny nesmí přesáhnout 12 hodin. Zároveň se zaměstnanec zavazuje, že při práci na dálku bude dodržovat příslušná ustanovení zákoníku práce upravující přestávky v práci a nepřetržitého denního a týdenního odpočinku.")
run.font.name = 'Arial'
# worker infor
worker = doc.add_paragraph()
worker_run = worker.add_run('Zaměstnanec: Ing. Filip Šrámek')
worker_run.font.name = 'Arial'
# Tabulka – 3 sloupce: Datum, časový interval, počet hodin
table = doc.add_table(rows=1, cols=3)
run.font.name = 'Arial'
run.font.size = Pt(11)

# ----- 3.1 Vyplnění hlavičky -----
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Datum'
hdr_cells[1].text = 'V čase od‑do (včetně uvedení času přestávky)'
hdr_cells[2].text = 'Počet hodin'

workedHours = 0

def duration_to_minutes(dur: str) -> float:
    """
    Převádí řetězec ve tvaru 'H:MM' nebo 'HH:MM' na počet minut.
    Příklad: '2:30' → 150
    """
    # Očistíme případné mezery
    dur = dur.strip()
    # Rozdělíme na hodiny a minuty
    parts = dur.split(':')
    if len(parts) != 2:
        raise ValueError(f"Neplatný formát trvání: {dur}")

    hours, minutes = parts

    intMinutes = int(minutes)
    if int(intMinutes) > 0:
        intMinutes = intMinutes / 60
    else:
        intMinutes = 0

    return int(hours) + float(intMinutes)

# ----- 3.2 Vyplnění datových řádků -----
for row in df.itertuples(index=False):
    row_cells = table.add_row().cells

    # row.Date je datetime → formátujeme do českého tvaru
    formatted_date = row.Date.strftime('%d.%m.%Y')

    values = [
        formatted_date,
        f"{row.From} - {row.To}",
        row.Duration
    ]

    workedHours += duration_to_minutes(row.Duration)

    for i, value in enumerate(values):
        row_cells[i].text = str(value)

row_cells = table.add_row().cells
row_cells[1].text = 'Počet hodin odpracovaných celkem/měsíc'
row_cells[2].text = f'{workedHours:.2f}'
# worker infor
worker_location = doc.add_paragraph()
worker_location_run = worker_location.add_run('V Praze dne 30.11.2025')
worker_location_run.font.name = 'Arial'

# worker infor
worker_sign = doc.add_paragraph()
worker_sign_run = worker_sign.add_run('Podpis zaměstnance: Filip Šrámek')
worker_sign_run.font.name = 'Arial'

# legal paragraph 2
paragraphtext = f'Potvrzení o převzetí dokončené práce na základě dohody o dohody o pracovní činnosti č. 1/2025 za měsíc {headingMonth} roku {df['Date'].dt.year.iloc[0]}.'

legalParagraph2 = doc.add_paragraph()
legalParagraph2_run = legalParagraph2.add_run(paragraphtext)
legalParagraph2_run.font.name = 'Arial'
legalParagraph2_run.bold = True

summaryParagraph1 = f'Práci, k níž se '
summaryParagraph2 = f'Ing. Filip Šrámek'
summaryParagraph3 = f' zavázal na základě dohody ze dne 11.11.2025, převzal odpovědný pracovník ÚJČ AV ČR, Mgr. Ondřej Svoboda, ve sjednaném rozsahu a kvalitě.Při provedení práce bylo odpracováno celkem {workedHours:.2f} hodin. K výplatě dle dohody náleží pracovníkovi částka ve výši {workedHours * HOURLY_RATE:.2f}Kč.'

summary_par = doc.add_paragraph()

# Text před ztučněním
run1 = summary_par.add_run(summaryParagraph1)
run1.font.name = 'Arial'

# **Ztučněná část** – např. počet odpracovaných hodin
run2 = summary_par.add_run(summaryParagraph2)
run2.font.name = 'Arial'
run2.bold = True                     # ← ztučnění

# Zbytek odstavce (neztučněný)
run3 = summary_par.add_run(summaryParagraph3)
run3.font.name = 'Arial'

# worker infor
worker_location = doc.add_paragraph()
worker_location_run = worker_location.add_run('V Praze dne 30.11.2025')
worker_location_run.font.name = 'Arial'

#controller worker sign
ControllerWorkerSign = doc.add_paragraph()
ControllerWorkerSign_run =ControllerWorkerSign.add_run('Podpis odpovědného pracovníka:')
ControllerWorkerSign_run.font.name = 'Arial'

#Check date and sign
check_date = doc.add_paragraph()
check_date_run = check_date.add_run('Kontrola a proplacení provedeno dne                                       podpis:  ')
check_date_run.font.name = 'Arial'
# -------------------------------------------------
# 4. Uložení dokumentu
# -------------------------------------------------
doc.save(f'FilipŠrámekExportMěsíc{headingMonth}.docx')