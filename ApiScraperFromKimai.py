import requests
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
import os, ssl, smtplib
from email.message import EmailMessage
from dotenv import load_dotenv, find_dotenv

load_dotenv(find_dotenv())

API_URL = os.getenv('KIMAI_API_URL', 'http://192.168.100.8:8001/api/timesheets')
ACTIVITIES_URL = os.getenv('KIMAI_ACTIVITIES_URL', 'http://192.168.100.8:8001/api/activities')
PROJECT_TITLE = os.getenv('KIMAI_PROJECT_TITLE', 'Lexikografický projekt 1')
BEARER_TOKEN = os.getenv('KIMAI_API_TOKEN')
OUTPUT_DIR = os.getenv('OUTPUT_DIR', '/app/output')
USERNAME = "admin"
HOURLY_RATE = 350  # Kč za hodinu
CZECH_MONTHS = {
    1:  'leden',
    2:  'únor',
    3:  'březen',
    4:  'duben',
    5:  'květen',
    6:  'červen',
    7:  'červenec',
    8:  'srpen',
    9:  'září',
    10: 'říjen',
    11: 'listopad',
    12: 'prosinec'
}


class ApiScraperFromKimai:
    def __init__(self, base_url: str, api_key: str, username: str):
        self.base_url = base_url
        self.api_key = api_key
        self.username = username
        self.headers = {"Authorization": f"Bearer {self.api_key}"}

    def get_active_activity_ids(self):
        response = requests.get(ACTIVITIES_URL, headers=self.headers, timeout=10)
        if response.status_code != 200:
            print("Chyba při stahování aktivit:", response.status_code, response.text)
            return []
        activities = response.json()
        ids = [a['id'] for a in activities if a.get('visible') is True and a.get('parentTitle') == PROJECT_TITLE]
        print(f"Aktivní activity IDs pro '{PROJECT_TITLE}': {ids}")
        return ids

    def get_timesheets(self):
        response = requests.get(self.base_url, headers=self.headers, timeout=10)
        if response.status_code != 200:
            print("Chyba při stahování dat:", response.status_code, response.text)
            return None
        return response.json()

    def process_timesheets(self):
        timesheets = self.get_timesheets()
        if timesheets is None:
            print("Nepodařilo se získat timesheety.")
            return None

        current_date = datetime.now()

        is_test = bool(os.getenv('SMTP_TO_OVERRIDE'))
        if is_test:
            first_day_previous = datetime(current_date.year, current_date.month, 1)
        elif current_date.month == 1:
            first_day_previous = datetime(current_date.year - 1, 12, 1)
        else:
            first_day_previous = datetime(current_date.year, current_date.month - 1, 1)

        active_ids = self.get_active_activity_ids()

        filtered = []
        for ts in timesheets:
            begin_str = ts["begin"]
            begin_dt = datetime.strptime(begin_str, "%Y-%m-%dT%H:%M:%S%z")
            activity_id = ts.get('activity', {}).get('id') if isinstance(ts.get('activity'), dict) else ts.get('activity')
            if first_day_previous.month == begin_dt.month and first_day_previous.year == begin_dt.year and ts['billable'] is True and activity_id in active_ids:
                filtered.append({
                    "Date": begin_dt,
                    "From": begin_dt.strftime("%H:%M"),
                    "To": ts["end"] and datetime.strptime(ts["end"], "%Y-%m-%dT%H:%M:%S%z").strftime("%H:%M") or "",
                    "Duration": str(int(ts["duration"] // 3600)) + ":" + str(int((ts["duration"] % 3600) // 60)).zfill(2)
                })

        print("Zpracování dokončeno.")
        return filtered


class DocumentGenerator:
    def __init__(self, data, hourly_rate):
        self.data = data
        self.hourly_rate = hourly_rate
        self.file_name = ""

    def generate_document(self):
        if self.data is None or len(self.data) == 0:
            print("Žádná data k vytvoření dokumentu.")
            return

        heading_month = CZECH_MONTHS[self.data[0]["Date"].month]
        doc_year = self.data[0]["Date"].year

        today = datetime.now()
        if os.getenv('SMTP_TO_OVERRIDE'):
            sign_date = today.strftime('%d.%m.%Y')
        else:
            first_of_this_month = today.replace(day=1)
            sign_date = (first_of_this_month - timedelta(days=1)).strftime('%d.%m.%Y')

        doc = Document()

        headingText = f'Výkaz odpracované doby zaměstnancem u dohody o pracovní činnosti za měsíc {heading_month} roku {doc_year}'
        heading = doc.add_heading(headingText, level=1)
        runHeading = heading.runs[0]
        runHeading.font.name = 'Arial'
        runHeading.font.size = Pt(16)
        runHeading.bold = True
        runHeading.font.color.rgb = RGBColor(0, 0, 0)

        intro = doc.add_paragraph()
        run = intro.add_run("Práce je po dohodě smluvních stran vykonávána na dálku. V souladu s § 87a zákoníku práce si bude zaměstnanec pracovní dobu do směn při práci na dálku rozvrhovat sám, práce nebude vykonávána ve dnech svátků, víkendů a v čase od 22:00 do 6:00 hod. Délka směny nesmí přesáhnout 12 hodin. Zároveň se zaměstnanec zavazuje, že při práci na dálku bude dodržovat příslušná ustanovení zákoníku práce upravující přestávky v práci a nepřetržitého denního a týdenního odpočinku.")
        run.font.name = 'Arial'
        run.font.size = Pt(11)

        worker = doc.add_paragraph()
        worker_run = worker.add_run('Zaměstnanec: Ing. Filip Šrámek')
        worker_run.font.name = 'Arial'

        table = doc.add_table(rows=1, cols=3)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Datum'
        hdr_cells[1].text = 'V čase od:do (včetně uvedení času přestávky)'
        hdr_cells[2].text = 'Počet hodin'

        workedHours = 0

        def duration_to_minutes(dur: str) -> float:
            dur = dur.strip()
            parts = dur.split(':')
            if len(parts) != 2:
                raise ValueError(f"Neplatný formát trvání: {dur}")
            hours, minutes = parts
            intMinutes = int(minutes)
            if int(intMinutes) > 0:
                intMinutes = intMinutes / 60
            else:
                intMinutes = 0
            return int(hours) + float(intMinutes)

        for row in self.data:
            row_cells = table.add_row().cells
            formatted_date = row["Date"].strftime('%d.%m.%Y')
            values = [
                formatted_date,
                f"{row['From']} - {row['To']}",
                duration_to_minutes(row['Duration'])
            ]
            workedHours += duration_to_minutes(row['Duration'])
            for i, value in enumerate(values):
                row_cells[i].text = str(value)

        row_cells = table.add_row().cells
        row_cells[1].text = 'Počet hodin odpracovaných celkem/měsíc'
        row_cells[2].text = f'{workedHours:.2f}'

        worker_location = doc.add_paragraph()
        worker_location.add_run(f'V Praze dne {sign_date}').font.name = 'Arial'

        worker_sign = doc.add_paragraph()
        worker_sign.add_run('Podpis zaměstnance: Filip Šrámek').font.name = 'Arial'

        paragraphtext = f'Potvrzení o převzetí dokončené práce na základě dohody o dohody o pracovní činnosti č. 1/2025 za měsíc {heading_month} roku {doc_year}.'
        legalParagraph2 = doc.add_paragraph()
        legalParagraph2_run = legalParagraph2.add_run(paragraphtext)
        legalParagraph2_run.font.name = 'Arial'
        legalParagraph2_run.bold = True

        summary_par = doc.add_paragraph()
        run1 = summary_par.add_run('Práci, k níž se ')
        run1.font.name = 'Arial'
        run2 = summary_par.add_run('Ing. Filip Šrámek')
        run2.font.name = 'Arial'
        run2.bold = True
        run3 = summary_par.add_run(f' zavázal na základě dohody ze dne 11.11.2025, převzal odpovědný pracovník ÚJČ AV ČR, Mgr. Ondřej Svoboda, ve sjednaném rozsahu a kvalitě.Při provedení práce bylo odpracováno celkem {workedHours:.2f} hodin. K výplatě dle dohody náleží pracovníkovi částka ve výši {workedHours * HOURLY_RATE:.2f}Kč.')
        run3.font.name = 'Arial'

        worker_location2 = doc.add_paragraph()
        worker_location2.add_run(f'V Praze dne {sign_date}').font.name = 'Arial'

        ControllerWorkerSign = doc.add_paragraph()
        ControllerWorkerSign.add_run('Podpis odpovědného pracovníka:').font.name = 'Arial'

        check_date = doc.add_paragraph()
        check_date.add_run('Kontrola a proplacení provedeno dne                                       podpis:  ').font.name = 'Arial'

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        file_name = f'Filip_Šrámek_Export_Měsíc_{heading_month}_{doc_year}.docx'
        self.file_name = os.path.join(OUTPUT_DIR, file_name)
        doc.save(self.file_name)
        print(f'Dokument uložen jako {self.file_name}')

    def get_file_name(self):
        return self.file_name


class SendEmail:
    def __init__(self, smtp_server, port, username, password):
        self.smtp_server = smtp_server
        self.port = port
        self.username = username
        self.password = password

    def send_email(self, attachment_path, recipient):
        msg = EmailMessage()
        msg['From'] = self.username
        msg['To'] = recipient
        msg['Subject'] = "Výkaz odpracované doby - Filip Šrámek"
        msg.set_content("Dobrý den,\n\nv příloze zasílám výkaz odpracované doby za uplynulý měsíc.\n\nS pozdravem,\nFilip Šrámek")

        with open(attachment_path, 'rb') as f:
            file_data = f.read()
        msg.add_attachment(
            file_data,
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(attachment_path),
        )
        context = ssl.create_default_context()
        try:
            with smtplib.SMTP(self.smtp_server, self.port, timeout=10) as smtp:
                smtp.ehlo()
                smtp.starttls(context=context)
                smtp.ehlo()
                smtp.login(self.username, self.password)
                smtp.send_message(msg)
                print("E-mail byl úspěšně odeslán.")
        except Exception as e:
            print("Chyba při odesílání e-mailu:", e)


if __name__ == "__main__":
    scraper = ApiScraperFromKimai(API_URL, BEARER_TOKEN, USERNAME)
    document = DocumentGenerator(scraper.process_timesheets(), HOURLY_RATE)
    document.generate_document()

    smtp_user = os.getenv('SMTP_USER')
    smtp_pass = os.getenv('SMTP_PASS')
    smtp_to = os.getenv('SMTP_TO_OVERRIDE') or os.getenv('SMTP_TO')
    if document.get_file_name() and smtp_user and smtp_pass and smtp_to:
        email = SendEmail("smtp.gmail.com", 587, smtp_user, smtp_pass)
        email.send_email(document.get_file_name(), smtp_to)
    elif not document.get_file_name():
        print("Dokument nebyl vytvořen — e-mail nebyl odeslán.")
    else:
        print("SMTP proměnné nejsou nastaveny — e-mail nebyl odeslán.")
